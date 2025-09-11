import asyncio
from playwright.async_api import async_playwright
from datetime import datetime, timedelta
from docx import Document
import time

async def crawl_global_hydrogen_news():
    # 创建Word文档
    doc = Document()
    doc.add_heading('全球氢能网新闻汇总', 0)
    doc.add_paragraph(f'报告生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph()
    
    async with async_playwright() as p:
        # 启动浏览器，设置更长的超时时间
        browser = await p.chromium.launch(headless=True)
        context = await browser.new_context()
        
        # 设置全局导航超时时间为60秒[1](@ref)
        context.set_default_navigation_timeout(60000)
        context.set_default_timeout(45000)
        
        page = await context.new_page()
        
        try:
            # 访问全球氢能网，使用更宽松的等待条件[3](@ref)
            print("正在访问全球氢能网...")
            await page.goto('http://www.china-nengyuan.com/news/', wait_until='domcontentloaded', timeout=60000)
            
            # 等待主要内容加载[6](@ref)
            try:
                await page.wait_for_selector('.news-list', timeout=45000)
            except Exception as e:
                print(f"等待选择器超时: {e}，尝试继续执行...")
            
            # 获取总页数（这里需要根据实际网站结构调整）
            total_pages = 2245  # 根据网站显示的总页数
            current_page = 1
            max_retries = 3
            news_count = 0
            
            while current_page <= total_pages and news_count < 100:  # 限制爬取数量用于测试
                print(f'正在爬取第 {current_page} 页...')
                
                # 重试机制[6](@ref)
                for attempt in range(max_retries):
                    try:
                        # 获取当前页的所有新闻项
                        news_items = await page.query_selector_all('.news-list .news-item')
                        if not news_items:
                            raise Exception("未找到新闻项")
                        
                        print(f"第 {current_page} 页找到 {len(news_items)} 条新闻")
                        break
                    except Exception as e:
                        if attempt == max_retries - 1:
                            print(f"获取第 {current_page} 页新闻失败: {e}")
                            break
                        print(f"获取新闻失败，第 {attempt+1} 次重试...")
                        await asyncio.sleep(2)
                
                # 处理当前页的新闻
                for item in news_items:
                    try:
                        # 提取新闻标题和链接
                        title_element = await item.query_selector('.news-title a')
                        if not title_element:
                            continue
                            
                        title = await title_element.inner_text()
                        link = await title_element.get_attribute('href')
                        if not link.startswith('http'):
                            link = f'http://www .china-nengyuan.com{link}'
                        
                        # 提取新闻日期
                        date_element = await item.query_selector('.news-date')
                        date_str = await date_element.inner_text() if date_element else "日期未知"
                        
                        # 将新闻添加到Word文档
                        doc.add_heading(title, level=2)
                        doc.add_paragraph(f'发布日期: {date_str}')
                        doc.add_paragraph(f'原文链接: {link}')
                        doc.add_paragraph('-' * 50)
                        
                        news_count += 1
                        
                    except Exception as e:
                        print(f'处理新闻时出错: {e}')
                        continue
                
                # 翻到下一页
                if current_page < total_pages:
                    next_page = current_page + 1
                    print(f"尝试翻到第 {next_page} 页...")
                    
                    # 尝试多种翻页方式
                    try:
                        # 方式1: 尝试点击下一页按钮
                        next_btn = await page.query_selector(f'a:has-text("{next_page}")')
                        if next_btn:
                            await next_btn.click()
                        else:
                            # 方式2: 尝试直接构造URL访问下一页
                            next_url = f'http://www.china-nengyuan.com/news/index.php?gopage={current_page}'
                            await page.goto(next_url, wait_until='domcontentloaded', timeout=60000)
                        
                        # 等待页面加载
                        await asyncio.sleep(2)
                        await page.wait_for_selector('.news-list', timeout=30000)
                        
                    except Exception as e:
                        print(f'翻页失败: {e}，尝试直接访问下一页URL')
                        try:
                            next_url = f'http://www.china-nengyuan.com/news/index.php?gopage={current_page}'
                            await page.goto(next_url, wait_until='domcontentloaded', timeout=60000)
                            await asyncio.sleep(2)
                        except Exception as retry_error:
                            print(f'直接访问下一页也失败: {retry_error}')
                            break
                
                current_page += 1
                # 添加延迟，避免请求过于频繁
                await asyncio.sleep(1)
            
            # 保存Word文档
            filename = f'全球氢能网新闻_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
            doc.save(filename)
            print(f'爬取完成，共获取 {news_count} 条新闻，结果已保存到: {filename}')
            
        except Exception as e:
            print(f'爬取过程中出错: {e}')
        finally:
            await browser.close()

if __name__ == '__main__':

    asyncio.run(crawl_global_hydrogen_news())