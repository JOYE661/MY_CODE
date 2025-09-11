import asyncio
from playwright.async_api import async_playwright
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import time
import re

# 创建Word文档
doc = Document()
# 添加一个标题
doc.add_heading('全球氢能网新闻', 0)

async def main():
    async with async_playwright() as p:
        # 启动浏览器，设置 headless=False 以便观察操作过程
        browser = await p.chromium.launch(headless=False)
        context = await browser.new_context()
        page = await context.new_page()
        
        # 设置默认导航超时时间为60秒
        page.set_default_navigation_timeout(60000)
        
        # 访问起始页面
        start_url = 'http://h2.china-nengyuan.com/news/index.php?gopage=0&'
        await page.goto(start_url)
        
        # 等待页面加载
        await page.wait_for_load_state('networkidle')
        
        # 获取总页数（这里需要根据实际网页结构调整选择器）
        # 假设页数通过某个元素显示，例如 ".total-pages"
        # 如果无法直接获取，可以设置一个最大页数或通过其他方式判断终止
        total_pages = 1  # 示例值，请根据实际情况调整或获取
        
        current_page = 0
        
        while current_page <= total_pages:
            print(f"正在处理第 {current_page} 页")
            
            # 构造当前页的URL
            page_url = f'http://h2.china-nengyuan.com/news/index.php?gopage={current_page}&'
            await page.goto(page_url)
            await page.wait_for_load_state('networkidle')
            
            # 获取当前页面的所有新闻链接列表
            # 需要根据实际网页结构调整选择器，这里使用示例选择器 '.news-list a'
            
            # 如果上述选择器无效，可以尝试更具体的选择器，或者使用 XPath
            # 例如：//div[@class='news-item']/a
            # news_links = page.get_by_role("link")
            
            # # 方法2: 如果上述选择器无效，尝试使用XPath定位
            # if len(news_links) == 0:
            #     news_links = await page.query_selector_all('//div[contains(@class, "news-list")]//li//a')
            
            # # 方法3: 如果仍然找不到，尝试查找包含特定文本或属性的链接
            # if len(news_links) == 0:
            #     news_links = await page.query_selector_all('a[href*="/news/"]')
            
            # print(f"第 {current_page} 页找到 {len(news_links)} 条新闻")
            news_links = await page.query_selector_all('//table//table//a[contains(@href, ".html")]')


            # 遍历当前页的每条新闻
            for i, link in enumerate(news_links):
                # 获取新闻链接的 href 属性
                href = await link.get_attribute('href')
                if href:
                    # 确保链接是绝对路径
                    if not href.startswith('http'):
                        href = f'http://h2.china-nengyuan.com{href}'
                    
                    # 在新标签页中打开新闻详情页
                    # async with context.expect_page() as new_page_info:
                    #     await link.click(modifiers=['Control', 'Shift'])  # 在新标签页打开
                        
                    # news_page = await new_page_info.value
                    news_page = await context.new_page()
                    await news_page.goto(href)  
                    
                    # 等待新页面加载
                    await news_page.wait_for_load_state('networkidle')
                    
                    # 获取新闻详情页的标题和内容
                    # 需要根据实际网页结构调整选择器
                    title_element = await news_page.query_selector('h1')
                    title = await title_element.inner_text() if title_element else "无标题"
                    
                    content_element = await news_page.query_selector('.content')
                    content = await content_element.inner_text() if content_element else "无内容"
                    
                    date_element = await news_page.query_selector('.date')
                    date = await date_element.inner_text() if date_element else "无日期"
                    
                    # 将信息添加到Word文档
                    doc.add_heading(title, level=1)
                    doc.add_paragraph(f"发布日期: {date}")
                    doc.add_paragraph(f"原文链接: {href}")
                    doc.add_paragraph("内容摘要:")
                    doc.add_paragraph(content[:500] + '...' if len(content) > 500 else content) # 只保存前500字符作为摘要
                    doc.add_paragraph('-' * 50)
                    
                    print(f"已抓取: {title}")
                    
                    # 关闭新闻详情页
                    await news_page.close()
                    
                    # 短暂的延迟，避免请求过于频繁
                    await asyncio.sleep(1)
            
            # 处理完当前页，尝试翻到下一页
            current_page += 1
            
            # 这里可以根据下一页按钮的存在与否来决定是否继续循环
            # 例如：next_button = await page.query_selector('text=下一页')
            # 如果找不到下一页按钮，可以 break 跳出循环
        
        # 保存Word文档
        doc.save('全球氢能网新闻.docx')
        print("所有新闻已保存到 Word 文档")
        
        # 关闭浏览器
        await browser.close()

# 运行异步主函数
asyncio.run(main())