from playwright.sync_api import sync_playwright, TimeoutError
from docx import Document
from docx.shared import Inches
from urllib.parse import urljoin
import time
import re

# -------------------------- 配置参数 --------------------------
# 目标列表页URL
BASE_LIST_URL = "http://h2.china-nengyuan.com/news/index.php?gopage=0&"
# 保存Word文档的文件夹（需提前创建，否则报错）
SAVE_DIR = "./hydrogen_news"
# 浏览器配置（无头模式/超时时间等）
BROWSER_CONFIG = {
    "headless": True,  # True=无界面运行，False=有界面调试
    "timeout": 30 * 1000  # 页面加载超时时间（毫秒）
}
# 反爬等待时间（单位：秒）
WAIT_TIME = 1.5


# -------------------------- 工具函数 --------------------------
def clean_filename(filename: str) -> str:
    """清理文件名中的特殊字符（避免保存失败）"""
    invalid_chars = r'[\/:*?"<>|]'
    return re.sub(invalid_chars, "_", filename.strip())


def save_to_word(news_data: dict, save_path: str):
    """将新闻数据保存为Word文档"""
    # 创建Word文档对象
    doc = Document()
    # 添加标题（加粗、居中）
    title_para = doc.add_heading(news_data["title"], level=1)
    title_para.alignment = 1  # 1=居中对齐
    # 添加发布时间
    doc.add_paragraph(f"发布时间：{news_data['publish_time']}", style="Intense Quote")
    # 添加正文
    doc.add_heading("正文内容：", level=2)
    for para in news_data["content"].split("\n"):
        if para.strip():  # 跳过空行
            doc.add_paragraph(para.strip())
    # 保存文档
    doc.save(save_path)
    print(f"✅ 已保存：{save_path}")


# -------------------------- 核心爬取逻辑 --------------------------
def crawl_news_detail(page, detail_url: str) -> dict:
    """爬取单条新闻详情页数据"""
    try:
        # 访问详情页
        page.goto(detail_url, timeout=BROWSER_CONFIG["timeout"])
        time.sleep(WAIT_TIME)  # 反爬等待

        # 提取新闻标题（根据实际页面结构调整选择器）
        # 说明：若页面标题在<title>标签中，可直接用page.title()；若在正文标题标签中，需替换选择器
        title = page.title().replace("- 中国能源网", "").strip()  # 去除标题后缀

        # 提取发布时间（示例选择器，需根据实际页面HTML调整）
        # 方法：打开详情页F12，找到发布时间元素的XPath/CSS选择器
        try:
            publish_time = page.locator("css=.publish-time").inner_text().strip()
        except:
            # 若无法提取具体时间，使用页面列表页的时间（需在列表页提前获取）
            publish_time = "未知时间"

        # 提取正文内容（示例选择器，需根据实际页面HTML调整）
        # 说明：通常正文在<div class="content">或<div id="article">等容器中
        content_elements = page.locator("css=.article-content p")  # 正文段落
        content = "\n\n".join([elem.inner_text().strip() for elem in content_elements.all()])

        # 若未提取到正文，尝试其他选择器（容错处理）
        if not content:
            content = page.locator("css=#main-content").inner_text().strip()

        return {
            "title": title,
            "publish_time": publish_time,
            "detail_url": detail_url,
            "content": content
        }

    except TimeoutError:
        print(f"❌ 详情页加载超时：{detail_url}")
        return None
    except Exception as e:
        print(f"❌ 爬取详情页失败：{detail_url}，错误：{str(e)}")
        return None


def crawl_news_list(page) -> tuple[list[str], bool]:
    """爬取列表页的新闻链接，返回（链接列表，是否有下一页）"""
    try:
        # 等待列表页加载完成（根据页面关键元素调整选择器）
        page.wait_for_selector("css=table tr", timeout=BROWSER_CONFIG["timeout"])
        time.sleep(WAIT_TIME)

        # 提取新闻链接（列表页中每个<tr>内的<a>标签）
        # 说明：使用urljoin拼接相对路径为完整URL
        link_elements = page.locator("css=table tr td a")  # 新闻标题链接
        news_links = []
        for elem in link_elements.all():
            href = elem.get_attribute("href")
            if href and "javascript" not in href:  # 过滤无效链接
                full_url = urljoin(BASE_LIST_URL, href)
                news_links.append(full_url)

        # 判断是否有下一页（根据翻页按钮的文本/属性调整）
        # 示例：下一页按钮通常为<a href="?gopage=1">下一页</a>或<button class="next">
        has_next_page = False
        try:
            # 检查"下一页"按钮是否存在且可点击
            next_btn = page.locator("css=a:has-text('下一页')")
            if next_btn.is_visible() and next_btn.is_enabled():
                has_next_page = True
        except:
            has_next_page = False

        return news_links, has_next_page

    except TimeoutError:
        print(f"❌ 列表页加载超时：{page.url}")
        return [], False
    except Exception as e:
        print(f"❌ 爬取列表页失败：{str(e)}")
        return [], False


# -------------------------- 主程序 --------------------------
def main():
    with sync_playwright() as p:
        # 启动浏览器（Chrome为例，可替换为firefox/edge）
        browser = p.chromium.launch(
            headless=BROWSER_CONFIG["headless"],
            args=["--no-sandbox"]  # 避免权限问题
        )
        context = browser.new_context(
            user_agent="Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/118.0.0.0 Safari/537.36"
        )
        page = context.new_page()

        try:
            # 1. 访问初始列表页
            page.goto(BASE_LIST_URL, timeout=BROWSER_CONFIG["timeout"])
            print(f"📌 已访问初始列表页：{BASE_LIST_URL}")

            page_num = 1  # 页码计数器
            while True:
                print(f"\n===== 开始爬取第 {page_num} 页 =====")

                # 2. 爬取当前列表页的新闻链接
                news_links, has_next = crawl_news_list(page)
                if not news_links:
                    print(f"⚠️  第 {page_num} 页未提取到新闻链接，跳过")
                    break

                print(f"📊 第 {page_num} 页共提取 {len(news_links)} 条新闻链接")

                # 3. 遍历链接，爬取详情并保存Word
                for idx, link in enumerate(news_links, 1):
                    print(f"\n🔍 正在爬取第 {page_num} 页第 {idx} 条：{link}")
                    news_data = crawl_news_detail(page, link)
                    if not news_data:
                        continue  # 爬取失败则跳过保存

                    # 生成Word文件名并保存
                    clean_title = clean_filename(news_data["title"])
                    save_path = f"{SAVE_DIR}/{page_num}_{idx}_{clean_title}.docx"
                    save_to_word(news_data, save_path)

                    # 退回列表页（继续下一条爬取）
                    page.go_back(timeout=BROWSER_CONFIG["timeout"])
                    time.sleep(WAIT_TIME)

                # 4. 处理翻页
                if has_next:
                    # 点击下一页按钮（根据实际页面选择器调整）
                    page.locator("css=a:has-text('下一页')").click()
                    page_num += 1
                    time.sleep(WAIT_TIME)
                else:
                    print(f"\n🎉 所有页面爬取完成，共爬取 {page_num} 页")
                    break

        except Exception as e:
            print(f"\n❌ 程序运行出错：{str(e)}")
        finally:
            # 关闭浏览器
            context.close()
            browser.close()


if __name__ == "__main__":
    main()