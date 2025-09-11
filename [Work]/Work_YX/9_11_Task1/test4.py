from playwright.sync_api import sync_playwright, TimeoutError
from docx import Document
from urllib.parse import urljoin
import time
import re

# -------------------------- 配置参数 --------------------------
TARGET_URL = "http://h2.china-nengyuan.com/news/"  # 目标首页
SAVE_DIR = "./hydrogen_news"  # Word保存目录（需提前创建）
BROWSER_CONFIG = {
    "headless": False,  # 调试阶段建议设为False（可视化操作）
    "timeout": 30 * 1000,
    "viewport": {"width": 1920, "height": 1080}  # 视图窗口大小（属于上下文配置）
}
WAIT_TIME = 2  # 操作间隔（反爬+页面加载）
SCROLL_STEP = 300  # 滚轮每次滚动的像素数
SCROLL_INTERVAL = 1  # 每次滚动后的等待时间（秒）


# -------------------------- 工具函数（保持不变） --------------------------
def clean_filename(filename: str) -> str:
    """清理文件名中的特殊字符（避免保存失败）"""
    invalid_chars = r'[\/:*?"<>|]'
    return re.sub(invalid_chars, "_", filename.strip())[:50]  # 限制文件名长度


def save_to_word(news_data: dict, save_path: str):
    """将新闻数据保存为Word文档"""
    doc = Document()
    # 标题（加粗居中）
    title_para = doc.add_heading(news_data["title"], level=1)
    title_para.alignment = 1
    # 发布时间
    doc.add_paragraph(f"发布时间：{news_data['publish_time']}", style="Intense Quote")
    # 正文
    doc.add_heading("正文内容：", level=2)
    for para in news_data["content"].split("\n"):
        if para.strip():
            doc.add_paragraph(para.strip())
    # 保存
    doc.save(save_path)
    print(f"✅ 已保存：{save_path}")


def is_scroll_to_bottom(page) -> bool:
    """判断是否已滚动到页面底部"""
    prev_height = page.evaluate("document.body.scrollHeight")
    page.mouse.wheel(0, SCROLL_STEP)
    time.sleep(SCROLL_INTERVAL)
    curr_height = page.evaluate("document.body.scrollHeight")
    return prev_height == curr_height


# -------------------------- 核心交互与爬取逻辑（保持不变） --------------------------
def simulate_middle_click_enter(page):
    """模拟鼠标点击屏幕中间位置进入页面（适配首页交互）"""
    print("📌 正在模拟鼠标点击屏幕中间进入页面...")
    # 获取屏幕中间坐标（基于viewport配置：1920x1080的中间是(960, 540)）
    middle_x = BROWSER_CONFIG["viewport"]["width"] // 2
    middle_y = BROWSER_CONFIG["viewport"]["height"] // 2

    page.mouse.move(middle_x, middle_y)
    time.sleep(1)
    page.mouse.click(middle_x, middle_y, delay=300)
    time.sleep(WAIT_TIME)

    if "news" not in page.url:
        raise Exception("❌ 鼠标点击未触发页面跳转，请检查首页中间是否为可点击元素")
    print(f"✅ 已通过中间点击进入目标页面：{page.url}")


def crawl_detail_page(page, detail_url: str) -> dict:
    """爬取详情页数据"""
    try:
        page.goto(detail_url, timeout=BROWSER_CONFIG["timeout"])
        time.sleep(WAIT_TIME)

        # 提取标题（根据实际页面调整选择器）
        title = page.locator("css=.article-title").inner_text().strip() if page.locator("css=.article-title").count() > 0 else page.title().strip()

        # 提取发布时间（示例选择器，需替换为页面真实元素）
        publish_time = page.locator("css=.publish-date").inner_text().strip() if page.locator("css=.publish-date").count() > 0 else "未知时间"

        # 提取正文（示例选择器：正文容器内的所有文本）
        content_container = page.locator("css=.article-content")
        if content_container.count() == 0:
            content_container = page.locator("css=#news-content")
        content = "\n\n".join([p.inner_text().strip() for p in content_container.locator("p").all() if p.inner_text().strip()])

        return {
            "title": title,
            "publish_time": publish_time,
            "detail_url": detail_url,
            "content": content
        }
    except TimeoutError:
        print(f"❌ 详情页加载超时：{detail_url}")
        return None
    except Exception as e:
        print(f"❌ 爬取详情页失败：{str(e)}")
        return None


def process_columns_by_scroll(page):
    """通过滚轮滑动遍历栏目，爬取每条新闻"""
    crawled_links = set()
    column_index = 1

    while True:
        print(f"\n===== 开始处理第 {column_index} 个栏目 =====")

        # 栏目选择器（需根据页面实际结构替换）
        column_selector = "css=.news-column"
        current_column = page.locator(column_selector).nth(column_index - 1)

        if current_column.count() == 0:
            print(f"⚠️  未找到第 {column_index} 个栏目，尝试滚动...")
            if is_scroll_to_bottom(page):
                print("📌 已滚动到页面底部，无更多栏目")
                break
            continue

        # 提取当前栏目的新闻链接
        news_links = []
        link_elements = current_column.locator("css=a.news-link")  # 需替换为实际链接选择器
        for elem in link_elements.all():
            href = elem.get_attribute("href")
            if href and href not in crawled_links and "javascript" not in href:
                full_url = urljoin(TARGET_URL, href)
                news_links.append(full_url)
                crawled_links.add(href)

        if not news_links:
            print(f"⚠️  第 {column_index} 个栏目无有效新闻链接")
            column_index += 1
            continue

        print(f"📊 第 {column_index} 个栏目共提取 {len(news_links)} 条新闻")

        # 爬取当前栏目的新闻详情
        for idx, link in enumerate(news_links, 1):
            print(f"\n🔍 爬取第 {column_index} 栏目第 {idx} 条：{link}")
            news_data = crawl_detail_page(page, link)
            if not news_data:
                continue

            # 保存为Word
            clean_title = clean_filename(news_data["title"])
            save_path = f"{SAVE_DIR}/栏目{column_index}_第{idx}_{clean_title}.docx"
            save_to_word(news_data, save_path)

            # 退回栏目列表页
            page.go_back(timeout=BROWSER_CONFIG["timeout"])
            time.sleep(WAIT_TIME)

        # 滚动到下一个栏目
        print(f"📜 第 {column_index} 个栏目处理完成，滚动到下一个栏目...")
        page.mouse.wheel(0, SCROLL_STEP * 2)
        time.sleep(SCROLL_INTERVAL * 2)
        column_index += 1

        if is_scroll_to_bottom(page):
            print("🎉 所有栏目处理完成！")
            break


# -------------------------- 主程序（修正viewport参数位置） --------------------------
def main():
    with sync_playwright() as p:
        # 启动浏览器（launch()方法中移除viewport参数）
        browser = p.chromium.launch(
            headless=BROWSER_CONFIG["headless"],  # headless属于launch的参数
            args=["--no-sandbox"]
            # 此处移除viewport，它不属于launch的参数
        )
        # 在上下文（new_context()）中配置viewport
        context = browser.new_context(
            user_agent="Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
            viewport=BROWSER_CONFIG["viewport"]  # 正确位置：viewport属于上下文参数
        )
        page = context.new_page()

        try:
            # 访问首页
            page.goto(TARGET_URL, timeout=BROWSER_CONFIG["timeout"])
            print(f"✅ 已访问首页：{TARGET_URL}")

            # 模拟鼠标中间点击进入新闻页面
            simulate_middle_click_enter(page)

            # 滚轮滑动遍历栏目，爬取新闻
            process_columns_by_scroll(page)

        except Exception as e:
            print(f"\n❌ 程序异常终止：{str(e)}")
        finally:
            # 关闭浏览器
            context.close()
            browser.close()


if __name__ == "__main__":
    main()
