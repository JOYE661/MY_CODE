from playwright.sync_api import sync_playwright, TimeoutError
from docx import Document
from urllib.parse import urljoin
import time
import re

# -------------------------- 配置参数 --------------------------
TARGET_URL = "http://h2.china-nengyuan.com/news/"  # 目标列表页
SAVE_DIR = "./hydrogen_news"  # Word保存目录（需提前创建）
BROWSER_CONFIG = {
    "headless": False,  # 调试时设为False（可视化），部署时改True
    "timeout": 30 * 1000,
    "viewport": {"width": 1920, "height": 1080}
}
WAIT_TIME = 2  # 操作间隔（反爬+页面加载）


# -------------------------- 工具函数 --------------------------
def clean_filename(filename: str) -> str:
    """清理文件名特殊字符"""
    invalid_chars = r'[\/:*?"<>|]'
    return re.sub(invalid_chars, "_", filename.strip())[:50]


def save_to_word(news_data: dict, save_path: str):
    """保存新闻数据为Word文档"""
    doc = Document()
    # 标题（加粗居中）
    title_para = doc.add_heading(news_data["title"], level=1)
    title_para.alignment = 1
    # 发布时间
    doc.add_paragraph(f"发布时间：{news_data['publish_time']}", style="Intense Quote")
    # 正文
    doc.add_heading("正文内容：", level=2)
    for para in news_data["content"].split("\n"):
        if para.strip():
            doc.add_paragraph(para.strip())
    doc.save(save_path)
    print(f"✅ 已保存：{save_path}")


# -------------------------- 详情页爬取 --------------------------
def crawl_detail_page(page, detail_url: str) -> dict:
    """爬取单条新闻详情页数据"""
    try:
        page.goto(detail_url, timeout=BROWSER_CONFIG["timeout"])
        time.sleep(WAIT_TIME)

        # 提取标题（优先取页面<h1>，否则用页面标题）
        title = page.locator("css=h1").inner_text().strip() if page.locator("css=h1").count() > 0 else page.title().strip()

        # 提取发布时间（示例选择器，需根据页面实际调整）
        publish_time = "未知时间"
        time_elems = page.locators("css=.publish-time, .date")  # 多个可能的时间选择器
        if time_elems.count() > 0:
            publish_time = time_elems.first.inner_text().strip()

        # 提取正文（示例选择器：正文容器内的所有<p>标签）
        content = ""
        content_containers = page.locators("css=.article-content, #news-content")  # 多个可能的正文容器
        if content_containers.count() > 0:
            content = "\n\n".join([
                p.inner_text().strip() 
                for p in content_containers.first.locator("p").all() 
                if p.inner_text().strip()
            ])

        return {
            "title": title,
            "publish_time": publish_time,
            "detail_url": detail_url,
            "content": content
        }
    except TimeoutError:
        print(f"❌ 详情页加载超时：{detail_url}")
        return None
    except Exception as e:
        print(f"❌ 爬取详情页失败：{str(e)}")
        return None


# -------------------------- 列表页链接提取 --------------------------
def extract_news_links(page) -> list[str]:
    """提取当前列表页所有“查看原文”的新闻链接"""
    try:
        # 等待“查看原文”链接加载（选择器：rel="nofollow"的<a>标签）
        page.wait_for_selector("css=a[rel='nofollow']", timeout=BROWSER_CONFIG["timeout"])
        time.sleep(WAIT_TIME)

        link_elements = page.locators("css=a[rel='nofollow']")
        news_links = []
        for elem in link_elements.all():
            href = elem.get_attribute("href")
            if href:
                full_url = urljoin(TARGET_URL, href)
                news_links.append(full_url)
        return news_links
    except TimeoutError:
        print("❌ 列表页链接加载超时")
        return []
    except Exception as e:
        print(f"❌ 提取列表页链接失败：{str(e)}")
        return []


# -------------------------- 主程序（翻页+爬取） --------------------------
def main():
    with sync_playwright() as p:
        # 启动浏览器
        browser = p.chromium.launch(
            headless=BROWSER_CONFIG["headless"],
            args=["--no-sandbox"]
        )
        context = browser.new_context(
            user_agent="Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
            viewport=BROWSER_CONFIG["viewport"]
        )
        page = context.new_page()

        try:
            current_page_url = TARGET_URL
            page_num = 1

            while True:
                print(f"\n===== 开始处理第 {page_num} 页：{current_page_url} =====")
                # 访问当前列表页
                page.goto(current_page_url, timeout=BROWSER_CONFIG["timeout"])
                time.sleep(WAIT_TIME)

                # 提取当前页所有新闻链接
                news_links = extract_news_links(page)
                if not news_links:
                    print("⚠️ 当前页无有效新闻链接，尝试翻页...")
                else:
                    print(f"📊 第 {page_num} 页共提取 {len(news_links)} 条新闻链接")

                    # 处理每条新闻
                    for idx, link in enumerate(news_links, 1):
                        print(f"\n🔍 处理第 {page_num} 页第 {idx} 条：{link}")
                        news_data = crawl_detail_page(page, link)
                        if not news_data:
                            continue

                        # 保存为Word
                        clean_title = clean_filename(news_data["title"])
                        save_path = f"{SAVE_DIR}/第{page_num}页_第{idx}条_{clean_title}.docx"
                        save_to_word(news_data, save_path)

                        # 返回列表页
                        page.go_back(timeout=BROWSER_CONFIG["timeout"])
                        time.sleep(WAIT_TIME)

                # 尝试翻页（下一页按钮选择器：包含“下一页”文字的<a>标签）
                try:
                    next_btn = page.locator("css=a:has-text('下一页')")
                    if next_btn.is_visible() and next_btn.is_enabled():
                        next_btn.click()
                        time.sleep(WAIT_TIME)
                        current_page_url = page.url
                        page_num += 1
                    else:
                        print("🎉 已无下一页，爬取结束！")
                        break
                except:
                    print("⚠️ 未找到下一页按钮或翻页失败，爬取结束！")
                    break

        except Exception as e:
            print(f"\n❌ 程序异常终止：{str(e)}")
        finally:
            # 关闭浏览器
            context.close()
            browser.close()


if __name__ == "__main__":
    main()