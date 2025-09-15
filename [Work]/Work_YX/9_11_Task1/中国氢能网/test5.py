import asyncio
from playwright.async_api import async_playwright
from docx import Document
import os
import re
import random
import pandas as pd
from datetime import datetime
import json
import aiohttp
import logging
import hashlib

# 设置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler("crawler.log"),
        logging.StreamHandler()
    ]
)

# 基础目录设置
script_dir = "/Users/joye/Desktop/MY_CODE/[Work]/Work_YX/9_11_Task1/中国氢能网/2025_News"
os.makedirs(script_dir, exist_ok=True)

# os.makedirs() - 递归创建目录的函数
# script_dir - 要创建的目录路径
# exist_ok=True - 如果目录已存在则不报错，避免重复创建时出现异常

# 创建数据目录结构
# 居然是个字典
data_dirs = {
    'news': os.path.join(script_dir, 'news_data'),
    'logs': os.path.join(script_dir, 'logs'),
    'proxies': os.path.join(script_dir, 'proxy_data')
}

for dir_name, dir_path in data_dirs.items():
    os.makedirs(dir_path, exist_ok=True)

# 用户代理列表
user_agents = [
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36",
    # 其他用户代理保持不变...
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36",
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/90.0.4430.212 Safari/537.36",
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/89.0.4389.114 Safari/537.36",
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/88.0.4324.150 Safari/537.36",
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/92.0.4515.107 Safari/537.36",
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/93.0.4577.63 Safari/537.36",
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/84.0.4147.125 Safari/537.36",
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/95.0.4638.54 Safari/537.36",
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/96.0.4664.45 Safari/537.36",
    "Mozilla/5.0 (Windows NT 6.3; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/77.0.3865.120 Safari/537.36",
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/81.0.4044.138 Safari/537.36",
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/82.0.4085.122 Safari/537.36",
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/79.0.3945.88 Safari/537.36",
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/78.0.3904.108 Safari/537.36",
    "Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/95.0.4638.54 Safari/537.36",
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/605.1.15 (KHTML, like Gecko) Chrome/91.0.4472.114 Safari/605.1.15",
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_6) AppleWebKit/605.1.15 (KHTML, like Gecko) Chrome/90.0.4430.212 Safari/605.1.15",
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_14_6) AppleWebKit/605.1.15 (KHTML, like Gecko) Chrome/89.0.4389.114 Safari/605.1.15",
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_14_5) AppleWebKit/605.1.15 (KHTML, like Gecko) Chrome/88.0.4324.150 Safari/605.1.15",
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_13_6) AppleWebKit/605.1.15 (KHTML, like Gecko) Chrome/87.0.4280.141 Safari/605.1.15",
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_5) AppleWebKit/605.1.15 (KHTML, like Gecko) Chrome/96.0.4664.45 Safari/605.1.15",
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_14_4) AppleWebKit/605.1.15 (KHTML, like Gecko) Chrome/84.0.4147.89 Safari/605.1.15",
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/86.0.4240.198 Safari/537.36",
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/83.0.4103.116 Safari/537.36",
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/82.0.4085.128 Safari/537.36",
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/81.0.4044.125 Safari/537.36",
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/80.0.3987.132 Safari/537.36",
    "Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/75.0.3770.90 Safari/537.36",
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/92.0.4515.107 Safari/537.36",
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/89.0.4389.110 Safari/537.36",
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/94.0.4606.61 Safari/537.36",
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/78.0.3904.97 Safari/537.36",
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/96.0.4664.45 Safari/537.36",
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_5) AppleWebKit/605.1.15 (KHTML, like Gecko) Chrome/82.0.4085.128 Safari/605.1.15",
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_6) AppleWebKit/605.1.15 (KHTML, like Gecko) Chrome/90.0.4430.212 Safari/605.1.15",
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/79.0.3945.130 Safari/537.36",
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/84.0.4147.135 Safari/537.36",
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/85.0.4183.121 Safari/537.36",
]

# 代理IP池
proxy_list = []

# 新闻日志文件
news_log_file = os.path.join(data_dirs['logs'], 'news_log.xlsx')

def load_existing_log():
    """加载已有的日志文件"""
    if os.path.exists(news_log_file):
        try:
            df = pd.read_excel(news_log_file)
            # 确保日期格式正确
            if '文档日期' in df.columns:
                df['文档日期'] = pd.to_datetime(df['文档日期'], errors='coerce').dt.strftime('%Y-%m-%d')
            return df
        except Exception as e:
            logging.error(f"加载日志文件失败: {e}")
            return pd.DataFrame(columns=news_log_columns)
    return pd.DataFrame(columns=news_log_columns)

# 新闻日志列定义
news_log_columns = ['文档ID', '文档名称', '文档链接', '是否有二级链接', '文档日期', '文件大小', '抓取时间', '状态']

# 加载现有日志
existing_log_df = load_existing_log()
existing_urls = set(existing_log_df['文档链接'].tolist()) if '文档链接' in existing_log_df.columns else set()
existing_titles = set(existing_log_df['文档名称'].tolist()) if '文档名称' in existing_log_df.columns else set()

# 创建一个新的DataFrame用于本次爬取
current_log_df = pd.DataFrame(columns=news_log_columns)

def generate_doc_id(url):
    """生成文档的唯一ID"""
    return hashlib.md5(url.encode()).hexdigest()[:12]

def update_news_log(news_info):
    """更新新闻日志"""
    global current_log_df, existing_log_df
    
    # 添加到当前日志
    current_log_df = pd.concat([current_log_df, pd.DataFrame([news_info])], ignore_index=True)
    
    # 定期保存日志（每8条保存一次）
    if len(current_log_df) % 8 == 0:
        save_combined_log()

def save_combined_log():
    """保存合并后的日志"""
    global current_log_df, existing_log_df
    
    # 合并新旧日志
    combined_df = pd.concat([existing_log_df, current_log_df], ignore_index=True)
    
    # 去重处理 - 保留最新的记录
    combined_df = combined_df.drop_duplicates(subset=['文档ID'], keep='last')
    
    try:
        combined_df.to_excel(news_log_file, index=False)
        logging.info(f"已保存 {len(combined_df)} 条日志记录")
    except Exception as e:
        logging.error(f"保存日志失败: {e}")

def sanitize_filename(filename):
    """清理文件名，移除非法字符"""
    return re.sub(r'[\\/*?:"<>|]', "_", filename)

def is_news_exists(news_url, news_title):
    """检查新闻是否已存在"""
    # 检查URL是否已存在
    if news_url in existing_urls:
        return True
    
    # 检查标题是否已存在（防止URL变化但内容相同的情况）
    if news_title in existing_titles:
        return True
    
    return False

async def fetch_proxies():
    """从外部API获取代理IP"""
    global proxy_list
    try:
        async with aiohttp.ClientSession() as session:
            # 这里可以替换为实际的代理API
            async with session.get('https://api.example.com/proxies') as response:
                if response.status == 200:
                    data = await response.json()
                    proxy_list = data.get('proxies', [])
                    
                    # 保存代理到文件
                    proxy_file = os.path.join(data_dirs['proxies'], 'proxies.json')
                    with open(proxy_file, 'w') as f:
                        json.dump(proxy_list, f)
                    
                    logging.info(f"从API获取了 {len(proxy_list)} 个代理")
    except Exception as e:
        logging.error(f"获取代理失败: {e}")
        # 如果获取失败，使用预定义的代理列表
        if not proxy_list:
            proxy_list = [
                # 备用代理...
            ]

async def get_random_proxy():
    """获取随机代理"""
    if not proxy_list:
        await fetch_proxies()
    
    if proxy_list:
        return random.choice(proxy_list)
    return None

async def create_browser_context(p, proxy_url=None):
    """创建浏览器上下文"""
    browser_args = {
        'headless': True,
        'timeout': 30000
    }
    
    # 如果有代理，设置代理
    if proxy_url:
        browser_args['proxy'] = {'server': proxy_url}
    
    browser = await p.chromium.launch(**browser_args)
    
    custom_user_agent = random.choice(user_agents)
    context = await browser.new_context(
        viewport={'width': 1920, 'height': 1080},
        user_agent=custom_user_agent,
        bypass_csp=True
    )
    
    # 拦截不必要的资源请求
    async def route_handler(route):
        if route.request.resource_type in ['image', 'stylesheet', 'font', 'media']:
            await route.abort()
        else:
            await route.continue_()
    
    await context.route('**/*', route_handler)
    return browser, context

async def process_news_link(context, link, page_url):
    """处理单个新闻链接"""
    news_info = {
        '文档ID': '',
        '文档名称': '',
        '文档链接': '',
        '是否有二级链接': 'NULL',
        '文档日期': '',
        '文件大小': 0,
        '抓取时间': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        '状态': '失败'
    }
    
    news_page = None
    try:
        # 获取新闻链接的 href 属性
        href = await link.get_attribute('href')
        if not href:
            return news_info
            
        # 确保链接是绝对路径
        if not href.startswith('http'):
            href = f'http://h2.china-nengyuan.com{href}'
        
        # 创建新闻ID
        doc_id = generate_doc_id(href)
        news_info['文档ID'] = doc_id
        news_info['文档链接'] = href
        
        # 创建新页面处理新闻详情
        news_page = await context.new_page()
        await news_page.goto(href, wait_until='domcontentloaded')
        
        # 获取新闻标题
        title_element = await news_page.query_selector('h1')
        title = await title_element.inner_text() if title_element else "无标题"
        news_info['文档名称'] = title
        
        # 检查是否已存在
        if is_news_exists(href, title):
            news_info['状态'] = '已存在'
            logging.info(f"跳过已存在新闻: {title}")
            await news_page.close()
            return news_info
        
        # 尝试获取发布日期
        date_element = await news_page.query_selector('.martop')  # 根据实际页面结构调整
        if date_element:
            date_text = await date_element.inner_text()
            # 从文本中提取日期
            date_match = re.search(r'\d{4}-\d{2}-\d{2}', date_text)
            if date_match:
                news_date = date_match.group(0)
                news_info['文档日期'] = news_date
        
        # 检查是否有二级链接
        secondary_links = await news_page.query_selector_all('//a[normalize-space()="查看原文"]')
        if len(secondary_links) > 0:  # 至少有一个除了当前链接外的其他链接
            url=await secondary_links[0].get_attribute('href') # 取第一个二级链接
            news_info['是否有二级链接'] =f"http://h2.china-nengyuan.com{url}"
        
        # 获取内容
        content_cell = await news_page.query_selector('.martop')
        full_content = await content_cell.inner_text() if content_cell else "无内容"
        
        # 创建文档
        doc = Document()
        doc.add_heading(title, level=1)
        doc.add_paragraph(f"原文链接: {href}")
        doc.add_paragraph(f"发布日期: {news_info['文档日期']}")
        doc.add_paragraph("内容:")
        doc.add_paragraph(full_content)
        
        # 清理标题用于文件名
        safe_title = sanitize_filename(title)
        
        # 按日期创建子目录
        date_dir = os.path.join(data_dirs['news'], news_info['文档日期'] if news_info['文档日期'] else 'unknown_date')
        os.makedirs(date_dir, exist_ok=True)
        
        # 保存文档
        save_path = os.path.join(date_dir, f"{safe_title}.docx")
        doc.save(save_path)
        
        # 记录文件大小
        file_size = os.path.getsize(save_path)
        news_info['文件大小'] = file_size
        news_info['状态'] = '成功'
        
        logging.info(f"已保存: {title} ({file_size} bytes)")
        
        # 关闭新闻详情页
        await news_page.close()
        
    except Exception as e:
        logging.error(f"处理新闻链接时出错: {e}")
        news_info['状态'] = f'错误: {str(e)}'
        # 确保页面被关闭
        if news_page:
            try:
                await news_page.close()
            except:
                pass
    
    return news_info

async def process_page(context, page_url, page_num):
    """处理单个页面"""
    logging.info(f"正在处理第 {page_num} 页: {page_url}")
    
    page = await context.new_page()
    await page.goto(page_url, wait_until='domcontentloaded')
    
    # 获取当前页面的所有新闻链接
    news_links = await page.query_selector_all('//table//table//h2/a[contains(@href, ".html")]')
    logging.info(f"第 {page_num} 页找到 {len(news_links)} 条新闻")
    
    # 处理所有新闻链接
    news_results = []
    for link in news_links:
        news_info = await process_news_link(context, link, page_url)
        news_results.append(news_info)
        update_news_log(news_info)
        
        # 添加随机延迟，避免请求过于频繁
        await asyncio.sleep(random.uniform(0.5, 2.0))
    
    # 检查是否有新数据，如果没有则提前停止
    new_news_count = sum(1 for info in news_results if info['状态'] == '成功')
    if new_news_count == 0 and page_num > 0:
        logging.info(f"第 {page_num} 页没有新数据，停止爬取")
        await page.close()
        return news_results, True  # 返回True表示应该停止爬取
    
    await page.close()
    return news_results, False

async def main():
    # 获取代理
    proxy_url = await get_random_proxy()
    logging.info(f"使用代理: {proxy_url if proxy_url else '无代理'}")
    
    async with async_playwright() as p:
        # 创建浏览器上下文
        browser, context = await create_browser_context(p, proxy_url)
        
        # 设置超时时间
        context.set_default_navigation_timeout(30000)
        context.set_default_timeout(20000)
        
        # 访问起始页面
        start_url = 'http://h2.china-nengyuan.com/news/index.php?gopage=0&'
        
        # 获取总页数（需要根据实际网站结构调整）
        total_pages = 3  # 示例值，实际应根据网站获取
        
        all_news_results = []
        stop_crawling = False
        
        for page_num in range(total_pages):
            if stop_crawling:
                logging.info("检测到无新数据，提前停止爬取")
                break
                
            page_url = f'http://h2.china-nengyuan.com/news/index.php?gopage={page_num}&'
            
            try:
                news_results, should_stop = await process_page(context, page_url, page_num)
                all_news_results.extend(news_results)
                stop_crawling = should_stop
                
                # 随机延迟 before next page
                await asyncio.sleep(random.uniform(1, 3))
                
            except Exception as e:
                logging.error(f"处理第 {page_num} 页时出错: {e}")
                # 可以在这里添加重试逻辑
        
        # 最终保存日志
        save_combined_log()
        
        # 统计结果
        new_count = sum(1 for info in all_news_results if info['状态'] == '成功')
        exist_count = sum(1 for info in all_news_results if info['状态'] == '已存在')
        error_count = sum(1 for info in all_news_results if info['状态'].startswith('错误'))
        
        logging.info(f"爬取完成: 新增 {new_count} 条, 已存在 {exist_count} 条, 错误 {error_count} 条")
        
        # 关闭浏览器
        await browser.close()

# 运行异步主函数
if __name__ == "__main__":
    asyncio.run(main())