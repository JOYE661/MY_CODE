import asyncio
from playwright.async_api import async_playwright
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import time
import re

async def main():
    # 创建Word文档
    doc = Document()
    # 添加一个标题
    doc.add_heading('全球氢能网新闻', 0)
    
    async with async_playwright() as p:
        # 启动浏览器，设置 headless=False 以便观察操作过程
        browser = await p.chromium.launch(headless=False)
        context = await browser.new_context()
        page = await context.new_page()
        
        # 设置默认导航超时时间为60秒
        page.set_default_navigation_timeout(60000)
        
        # 访问起始页面
        start_url = 'http://h2.china-nengyuan.com/news/'
        await page.goto(start_url)
        
        # 等待页面加载
        await page.wait_for_load_state('networkidle')
        
        # 获取总页数
        # 通过分析页面结构，找到页数信息
        page_info = await page.query_selector('div.page')
        page_text = await page_info.inner_text() if page_info else ""
        
        # 从文本中提取总页数
        total_pages_match = re.search(r'/(\d+)', page_text)
        total_pages = int(total_pages_match.group(1)) if total_pages_match else 3
        
        print(f"总页数: {total_pages}")
        
        current_page = 1
        
        while current_page <= total_pages:
            print(f"正在处理第 {current_page} 页")
            
            # 如果当前不是第一页，需要导航到指定页码
            if current_page > 1:
                page_url = f'http://h2.china-nengyuan.com/news/index.php?gopage={current_page-1}'
                await page.goto(page_url)
                await page.wait_for_load_state('networkidle')
            
            # 获取当前页面的所有新闻链接

            news_links = await page.query_selector_all('//table//table//a[contains(@href, ".html")]')
            # /html/body/table[3]/tbody/tr/td[1]/table/tbody/tr/td/table[5]/tbody/tr[1]/td[1]/h2/a
            
            print(f"第 {current_page} 页找到 {len(news_links)} 条新闻")
            
            # 遍历当前页的每条新闻
            for i, link in enumerate(news_links):
                try:
                    # 获取新闻链接的 href 属性
                    href = await link.get_attribute('href')
                    if href:
                        # 确保链接是绝对路径
                        if not href.startswith('http'):
                            href = f'http://h2.china-nengyuan.com{href}'
                        
                        # 创建新页面
                        news_page = await browser.new_page()
                        
                        try:
                            # 在新页面中访问新闻链接
                            await news_page.goto(href)
                            await news_page.wait_for_load_state('networkidle')
                            
                            # 获取新闻详情页的标题和内容
                            title_element = await news_page.query_selector('h1')
                            title = await title_element.inner_text() if title_element else "无标题"
                            
                            # 尝试多种可能的内容选择器
                            content_element = await news_page.query_selector('.content, .article-content, #content')
                            if not content_element:
                                # 如果上述选择器找不到，尝试通过XPath查找
                                content_element = await news_page.query_selector('//div[contains(@class, "content") or contains(@class, "article")]')
                            
                            content = await content_element.inner_text() if content_element else "无内容"
                            
                            # 尝试查找日期信息
                            date_element = await news_page.query_selector('.date, .time, .publish-date')
                            date = await date_element.inner_text() if date_element else "无日期"
                            
                            # 将信息添加到Word文档
                            doc.add_heading(title, level=1)
                            doc.add_paragraph(f"发布日期: {date}")
                            doc.add_paragraph(f"原文链接: {href}")
                            doc.add_paragraph("内容摘要:")
                            # 只保存前500字符作为摘要
                            summary = content[:500] + '...' if len(content) > 500 else content
                            doc.add_paragraph(summary)
                            doc.add_paragraph('-' * 50)
                            
                            print(f"已抓取: {title}")
                            
                        except Exception as e:
                            print(f"处理新闻时出错: {e}")
                        finally:
                            # 关闭新闻详情页
                            await news_page.close()
                        
                        # 短暂的延迟，避免请求过于频繁
                        await asyncio.sleep(1)
                
                except Exception as e:
                    print(f"处理链接时出错: {e}")
            
            # 处理完当前页，翻到下一页
            current_page += 1
        
        # 保存Word文档
        doc.save('全球氢能网新闻.docx')
        print("所有新闻已保存到 Word 文档")
        
        # 关闭浏览器
        await browser.close()

# 运行异步主函数
asyncio.run(main())