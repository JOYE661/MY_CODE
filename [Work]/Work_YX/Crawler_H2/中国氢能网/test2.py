import asyncio
from playwright.async_api import async_playwright
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import time
import re

# 创建Word文档
doc = Document()
# 添加一个标题
doc.add_heading('全球氢能网新闻', 0)

async def main():
    async with async_playwright() as p:
        # 启动浏览器，设置 headless=True 提高速度
        browser = await p.chromium.launch(headless=True)
        
        # 创建上下文时禁用图片加载
        context = await browser.new_context(
            viewport={'width': 1920, 'height': 1080},
            user_agent='Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
            # 跳过图片和CSS加载
            bypass_csp=True
        )
        
        # 拦截不必要的资源请求（图片、样式表、字体等）
        async def route_handler(route):
            if route.request.resource_type in ['image', 'stylesheet', 'font', 'media']:
                await route.abort()
            else:
                await route.continue_()
        
        await context.route('**/*', route_handler)
        
        page = await context.new_page()
        
        # 设置默认导航超时时间为30秒（减少等待时间）
        page.set_default_navigation_timeout(30000)
        page.set_default_timeout(20000)
        
        # 访问起始页面
        start_url = 'http://h2.china-nengyuan.com/news/index.php?gopage=0&'
        await page.goto(start_url, wait_until='domcontentloaded')  # 只等待DOM加载完成
        
        # 获取总页数
        total_pages = 2  # 示例值，请根据实际情况调整或获取
        
        current_page = 0
        
        while current_page <= total_pages:
            print(f"正在处理第 {current_page} 页")
            
            # 构造当前页的URL
            page_url = f'http://h2.china-nengyuan.com/news/index.php?gopage={current_page}&'
            await page.goto(page_url, wait_until='domcontentloaded')
            
            # 获取当前页面的所有新闻链接列表
            news_links = await page.query_selector_all('//table//table//h2/a[contains(@href, ".html")]')
            print(f"第 {current_page} 页找到 {len(news_links)} 条新闻")
            
            # 使用异步任务列表处理所有新闻链接
            tasks = []
            for link in news_links:
                tasks.append(process_news_link(context, link, doc))
            
            # 并发处理所有新闻链接，限制并发数为5避免过多请求
            for i in range(0, len(tasks), 5):
                batch = tasks[i:i+5]
                await asyncio.gather(*batch)
                await asyncio.sleep(1)  # 批次之间短暂延迟
            
            # 处理完当前页，尝试翻到下一页
            current_page += 1
            
            # 这里可以根据下一页按钮的存在与否来决定是否继续循环
            # 例如：next_button = await page.query_selector('text=下一页')
            # 如果找不到下一页按钮，可以 break 跳出循环
        
        # 保存Word文档
        doc.save('全球氢能网新闻.docx')
        print("所有新闻已保存到 Word 文档")
        
        # 关闭浏览器
        await browser.close()

async def process_news_link(context, link, doc):
    try:
        # 获取新闻链接的 href 属性
        href = await link.get_attribute('href')
        if not href:
            return
            
        # 确保链接是绝对路径
        if not href.startswith('http'):
            href = f'http://h2.china-nengyuan.com{href}'
        
        # 创建新页面处理新闻详情
        news_page = await context.new_page()
        await news_page.goto(href, wait_until='domcontentloaded')
        
        # 获取新闻详情页的标题和内容
        title_element = await news_page.query_selector('h1')
        title = await title_element.inner_text() if title_element else "无标题"
        
        content_cell = await news_page.query_selector('.martop')
        full_content = await content_cell.inner_text() if content_cell else "无内容"
        
        # 将信息添加到Word文档（注意：多线程操作文档需要加锁，这里简化处理）
        doc.add_heading(title, level=1)
        doc.add_paragraph(f"原文链接: {href}")
        doc.add_paragraph("内容摘要:")
        doc.add_paragraph(full_content[:1000] + '...' if len(full_content) > 1000 else full_content)
        doc.add_paragraph('-' * 50)
        
        print(f"已抓取: {title}")
        
        # 关闭新闻详情页
        await news_page.close()
        
    except Exception as e:
        print(f"处理新闻链接时出错: {e}")
        # 确保页面被关闭
        try:
            await news_page.close()
        except:
            pass

# 运行异步主函数
asyncio.run(main())